import json
import time
import uuid
from io import BytesIO

import PyPDF2
import pandas as pd
import requests
# get Scrap base url from settings
from django.conf import settings \
    # scrap base url
from docx import Document
# from celery import shared_task
from dualityAI.celery import app
from helpers.mongodb_client_data import get_details_by_id, update_analyzation, \
    update_analytic, get_all_session_by_subdomain
from helpers.openAI import analyze_article, analyze_conversation as analyze_conversation_openai
from requests_toolbelt.multipart.encoder import MultipartEncoder

from .models import ScrapeLink, Sitemap, Chatbot

scrap_base_url = settings.SCRAP_BASE_URL


# function to fetch data from ur curl -X 'POST' \
#   'https://scrap.duality.seethos.com/scrape-sitemap' \
#   -H 'accept: application/json' \
#   -H 'Content-Type: application/json' \
#   -d '{
#   "url": "https://morescript.com/wp-sitemap.xml"
# }'
def get_data_from_url(list_of_url=[]):
    base_url = f"{scrap_base_url}/scrapper/scrape-sitemaps"
    status_url = "https://scrap.duality.seethos.com/get-status"
    get_data_url = "https://scrap.duality.seethos.com/get-text"
    headers = {
        'accept': 'application/json',
        'Content-Type': 'application/json',
    }
    data = []

    # check if status of urls in list is scrape or not
    # if scrape then remove from list
    for url in list_of_url:
        # get scrape link by url
        scrape_link = ScrapeLink.objects.filter(url=url['url'], site_map=url['site_map']).first()
        # check if scrape link is  none or is_scrape is false if not none
        try:
            if scrape_link == None:
                # remove url from list
                data.append(url['url'])
            elif scrape_link.is_scrape == False:
                # remove url from list
                data.append(url['url'])
            else:
                pass
        except:
            pass

    # send data to url
    # limit data to 5
    data = data[:5]
    print(data)

    response = requests.post(base_url, headers=headers, data=json.dumps(data))

    # its response is {"message": "Scraping started in the background.","request_id": "9f0f31c3-085c-42f1-a055-707f59d1906e","status_url": "http://localhost:8000/get-status?request_id=9f0f31c3-085c-42f1-a055-707f59d1906e"}
    # we nee to get request_id and check status_url
    # get request_id
    print(response.status_code)
    print(response.json())
    request_id = response.json()['request_id']

    status = requests.get(f"{status_url}?request_id={request_id}")

    # check status
    if status.status_code == 200:
        status = status.json()['status']
        print(status)
        # status is "in_progress" wait until it is completed and run this in background
        while status == "in_progress":
            status = requests.get(f"{status_url}?request_id={request_id}")
            print(status)
            status = status.json()['status']
            # sleep for 5 seconds
            print(status)
            time.sleep(10)

    # check status if it is completed
    if status == "complete":
        response = requests.get(f"{get_data_url}?request_id={request_id}")

        # Check the response
        if response.status_code == 200:
            print('Request successful')
            # print(response.text)
            # set scrape status to true for all url i.e  list
            for url in list_of_url:
                # get scrape link by url
                scrape_link = ScrapeLink.objects.filter(url=url['url'], site_map=url['site_map']).first()
                # check if scrape link is  none or is_scrape is false if not none
                try:
                    if scrape_link == None:
                        # create scrape link
                        ScrapeLink.objects.create(url=url['url'], site_map=url['site_map'], is_scrape=True)
                    elif scrape_link.is_scrape == False:
                        # update scrape link
                        scrape_link.is_scrape = True
                        scrape_link.save()
                    else:
                        pass
                except:
                    pass

            return response.text
        else:
            print(f'Request failed with status code: {response.status_code}')
            print(response.text)
            return {"error": "Failed to get data from url"}

    else:
        print(f'Request failed with status code: {response.status_code}')
        print(response.text)
        return {"error": "Failed to get data from url"}


import uuid

@app.task
def create_chat_bot(chatbot, avatar_url="", file_paths=[]):
    print("DEBUG: Celery task create_chat_bot started.")
    print(f"DEBUG: create_chat_bot received chatbot ID: {chatbot}, avatar_url: {avatar_url}, file_paths: {file_paths}")
    # get chatbot
    chatbot = Chatbot.objects.filter(id=chatbot).first()
    if not chatbot:
        print(f"Error: Chatbot with ID {chatbot} not found in create_chat_bot task.")
        return {"error": "Chatbot not found"}
    # get data from chatbot
    name = chatbot.name
    questions = chatbot.question

    calendly = chatbot.call_to_action_link
    company_name = chatbot.company_name
    bot_id = chatbot.bot_id
    # Generate bot_id if it's not already set
    if not bot_id:
        bot_id = "bot" + str(uuid.uuid4())[0:8]
        chatbot.bot_id = bot_id
        chatbot.save()
    bot_id = bot_id.lower()
    
    container_id = chatbot.container_id
    language = chatbot.language
    color = chatbot.color

    avatar = chatbot.avatar
    conversation_start_flow = chatbot.conversation_start_flow
    conversation_end_flow = chatbot.conversation_end_flow

    greeting_message = chatbot.greeting_message
    chatbot_id = chatbot.id

    call_to_action_prompt = chatbot.call_to_action_prompt
    greeting_tags = chatbot.greeting_tags
    subscription_type = chatbot.account.subscription_type

    theme = chatbot.theme

    avatar_url = avatar_url
    agent_role = chatbot.agent_role
    custom_prompt = chatbot.custom_prompt

    # name='bluee', questions='name , email', file_path='test', calendly="http://test.com", company_name="seethos", bot_id="",container_id="",language='english',color='blue',avatar='test',conversation_start_flow='test',conversation_end_flow='test',greeting_message='Hello',chatbot_id="1",sitemap_url="",call_to_action_prompt="",greeting_tags=["Hello"],subscription_type="LOW",avatar_url="",theme="dark",chatbot=None):

    # Download the actual PDF file from the S3 URL
    pdf_binary_content = None
    txt_content = "" # This will accumulate text from all supported files

    # Process sitemap URLs
    sitemap_url = Sitemap.objects.filter(chatbot=chatbot).all()
    if sitemap_url:
        links = []
        for url in sitemap_url:
            url_list = url.traning_links
            for link in url_list:
                links.append({'url': link, 'site_map': url.id})

        data = get_data_from_url(links)
        try:
            chatbot.website_text += data
        except:
            chatbot.website_text = data
        chatbot.save()
        try:
            txt_content += chatbot.website_text
        except:
            pass

    # Process file_paths
    for file in file_paths:
        file_url = file[0]
        file_type = file[1]
        response = requests.get(file_url)

        if response.status_code != 200:
            print(f"Failed to download the file from S3: {file_url}")
            continue

        if file_type == "pdf":
            pdf_binary_content = response.content # Store binary content for PDF
            # Extract text from PDF for combined text content
            pdf_text = ""
            with BytesIO(response.content) as pdf_file_buffer:
                pdf_reader = PyPDF2.PdfReader(pdf_file_buffer)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    pdf_text += page.extract_text()
            txt_content += pdf_text
        elif file_type == "excel":
            data = BytesIO(response.content)
            xls = pd.read_excel(data, sheet_name=None)
            for sheet_name, df in xls.items():
                txt_content += df.to_string()
        elif file_type == "word":
            docx_bytes_io = BytesIO(response.content)
            doc = Document(docx_bytes_io)
            doc_text = ""
            for paragraph in doc.paragraphs:
                doc_text += paragraph.text + "\n"
            txt_content += doc_text
        elif file_type == "txt":
            txt_content += response.text
        else:
            print(f"File type {file_type} is not supported for text extraction.")

    # Prepare fields for MultipartEncoder
    fields = {
        "questions": json.dumps(questions),
        'greeting_tags': json.dumps(greeting_tags.split(",") if isinstance(greeting_tags, str) else greeting_tags),
        'avatar': ('avatar.txt', txt_content, 'text/plain'), # Still sending text for avatar
        'name': name,
        'calendly': calendly,
        'company_name': company_name,
        'language': language,
        'color': color,
        'conversation_start_flow': conversation_start_flow,
        'conversation_end_flow': conversation_end_flow,
        "greeting_message": greeting_message,
        "call_to_action_prompt": call_to_action_prompt,
        "subscription_type": subscription_type,
        "avatar_url": avatar_url,
        "theme": theme,
        "assistant_role": agent_role,
        "bot_id": bot_id,
        "system_prompt": custom_prompt
    }

    # Add training_data (accumulated text content) to the payload
    fields['training_data'] = ('training_data.txt', txt_content, 'text/plain')

    payload = MultipartEncoder(fields=fields)

    headers = {
        'Content-Type': payload.content_type
    }

    # Make the POST request
    url = f'{settings.INTG_BASE_URL}/bot/create_or_update/'  # Corrected API endpoint for seethos-bot-manager

    # greeting_tags = str(greeting_tags)
    query_parameters = {
        'name': name,
        'calendly': calendly,
        'company_name': company_name,
        'language': language,
        'color': color,
        'conversation_start_flow': conversation_start_flow,
        'conversation_end_flow': conversation_end_flow,
        "calendy": calendly,
        "greeting_message": greeting_message,
        "call_to_action_prompt": call_to_action_prompt,
        "greeting_tags": greeting_tags,
        "subscription_type": subscription_type,
        "avatar_url": avatar_url,
        "theme": theme,
        "assistant_role": agent_role,
        "bot_id": bot_id,
        "system_prompt": custom_prompt
    }

    print(f"DEBUG: Sending request to integration service: {url}")
    print(f"DEBUG: Request Headers: {headers}")
    # Note: Printing payload directly might be very verbose for large files.
    # For debugging, you might want to inspect `payload` object or its parts.
    # print(f"DEBUG: Request Payload (first 500 chars): {str(payload)[:500]}")
    response = requests.post(url, data=payload, headers=headers)
    print(f"DEBUG: Integration service response status code: {response.status_code}")
    print(f"DEBUG: Integration service response content: {response.text}")

    # Check the response
    if response.status_code == 200:
        print("Request successful, processing response.")
        bot = response.json()
        try:
            chatbot.bot_id = bot['subdomain']
            chatbot.status = True
            chatbot.traning_status = 'TRAINED'
            chatbot.save()
            print("Chatbot status updated successfully.")
        except Exception as e:
            print(f"Error updating chatbot status: {e}")
            pass
        print("Celery task create_chat_bot finished successfully.")
        return response.json()
    else:
        print(f"Request failed with status code: {response.status_code}")
        print(f"Integration service response content: {response.text}")
        print("Celery task create_chat_bot finished with error.")
        return {"error": "Failed to create chatbot"}


@app.task
def update_chat_bot(chatbot, avatar_url="", file_paths=[]):
    # get chatbot
    chatbot = Chatbot.objects.filter(id=chatbot).first()
    # get data from chatbot
    name = chatbot.name
    questions = chatbot.question
    calendly = chatbot.call_to_action_link
    company_name = chatbot.company_name
    bot_id = chatbot.bot_id.lower()
    container_id = chatbot.container_id
    language = chatbot.language
    color = chatbot.color
    avatar = chatbot.avatar
    conversation_start_flow = chatbot.conversation_start_flow
    conversation_end_flow = chatbot.conversation_end_flow
    greeting_message = chatbot.greeting_message
    chatbot_id = chatbot.id
    call_to_action_prompt = chatbot.call_to_action_prompt
    greeting_tags = chatbot.greeting_tags
    subscription_type = chatbot.account.subscription_type

    agent_role = chatbot.agent_role
    custom_prompt = chatbot.custom_prompt

    theme = chatbot.theme
    if not bot_id:
        # generate bot id
        bot_id = "bot" + str(uuid.uuid4())[0:8]
    avatar_url = avatar_url

    # name='bluee', questions='name , email', file_path='test', calendly="http://test.com", company_name="seethos", bot_id="",container_id="",language='english',color='blue',avatar='test',conversation_start_flow='test',conversation_end_flow='test',greeting_message='Hello',chatbot_id="1",sitemap_url="",call_to_action_prompt="",greeting_tags=["Hello"],subscription_type="LOW",avatar_url="",theme="dark",chatbot=None):

    # Download the actual PDF file from the S3 URL
    pdf_content = ""
    excel_content = ""
    word_content = ""
    txt_content = ""
    # file_paths = File.objects.filter(chatbot=chatbot_id).all()
    # get data from url
    # chatbot = Chatbot.objects.filter(id=chatbot).first()
    sitemap_url = Sitemap.objects.filter(chatbot=chatbot).all()
    if sitemap_url:
        # url = sitemap_url
        # get all url from sitemap and make a list
        links = []
        for url in sitemap_url:
            url_list = url.traning_links
            for link in url_list:
                links.append({'url': link, 'site_map': url.id})

        # print(links)
        text = get_data_from_url(links)
        # get website text from data site map and att data to it
        try:
            chatbot.website_text += text
        except:
            chatbot.website_text = text
        chatbot.save()
        try:
            txt_content += chatbot.website_text
        except:
            pass
    print(file_paths)
    for file in file_paths:
        file_path = file[0]
        file_type = file[1]
        response = requests.get(file_path)

        if response.status_code != 200:
            print("Failed to download the file from S3.")
            return

        if file_type == "pdf":
            # Extract text from PDF
            pdf_text = ""
            with BytesIO(response.content) as pdf_file_buffer:
                pdf_reader = PyPDF2.PdfReader(pdf_file_buffer)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    pdf_text += page.extract_text()

            # Append the extracted text to txt_content
            txt_content += pdf_text
            # print(txt_content)
        elif file_type == "excel":
            # extraxt text from excel
            data = BytesIO(response.content)

            # Read the file
            xls = pd.read_excel(data, sheet_name=None)

            # Print the content of each sheet
            for sheet_name, df in xls.items():
                # print(f"Sheet name: {sheet_name}")
                txt_content += df.to_string()

        elif file_type == "word":
            # Extract text from Word
            word_text = response.content

            # Load the DOCX document
            docx_bytes_io = BytesIO(response.content)
            doc = Document(docx_bytes_io)

            # Initialize an empty string to store the text from the document
            doc_text = ""

            # Iterate through the paragraphs in the document and extract the text
            for paragraph in doc.paragraphs:
                doc_text += paragraph.text + "\n"
            txt_content += doc_text

        else:
            print("File type is not supported")

    form_data = {
        'questions': questions,
        'greeting_tags': str(greeting_tags),

    }

    # Create a MultipartEncoder for the request data
    # check if greeting tags is string or list
    if type(greeting_tags) == str:
        greeting_tags = greeting_tags.split(",")

    print(txt_content)
    print("traning content\n", txt_content)
    payload = MultipartEncoder(
        fields={
            "questions": json.dumps(questions),
            'greeting_tags': json.dumps(greeting_tags),
            'pdf_file': ('txt_file.txt', txt_content, 'text/plain'),
            'avatar': ('avatar.txt', txt_content, 'text/plain')
        }
    )
    # print(payload)
    headers = {
        'Content-Type': payload.content_type
    }

    # Make the POST request

    # Make the POST request
    url = f"{settings.INTG_BASE_URL}/update_container/{bot_id}"
    # greeting_tags = str(greeting_tags)
    # print(greeting_tags)

    query_parameters = {
        'name': name,
        'calendly': calendly,
        'company_name': company_name,
        'bot_id': bot_id,
        'language': language,
        'color': color,
        'conversation_start_flow': conversation_start_flow,
        'conversation_end_flow': conversation_end_flow,
        "calendy": calendly,
        "greeting_message": greeting_message,
        "call_to_action_prompt": call_to_action_prompt,
        "greeting_tags": greeting_tags,
        "subscription_type": subscription_type,
        "avatar_url": avatar_url,
        "theme": theme,
        "assistant_role": agent_role,
        "system_prompt": custom_prompt
    }
    print("file upload test")
    response = requests.post(url, data=payload, headers=headers, json=query_parameters)
    print(response)
    try:
        print(response.json())
    except:
        print(response.text)

    # Check the response
    if response.status_code == 200:
        print('Request successful')
        bot = response.json()
        try:
            # chatbot.bot_id = bot['subdomain']
            #chatbot.container_id = bot['container_id']
            chatbot.status = True
            chatbot.traning_status = 'TRAINED'
            try:
                chatbot.bot_id = bot['subdomain']
            except:
                pass
            chatbot.save()
        except:
            pass
        # return response text
        return response.json()
    else:
        print(f'Request failed with status code: {response.status_code}')
        print(response.text)
        return {"error": "Failed to update chatbot"}

    # delete chatbot


def delete_chat_bot(container_id=""):
    url = f'{settings.INTG_BASE_URL}/delete_container/{container_id}'

    # request delete method
    response = requests.delete(url)
    # check response
    if response.status_code == 200:
        print('Request successful')
        return response.json()
    else:
        print(f'Request failed with status code: {response.status_code}')
        print(response.text)
        return {"error": "Failed to delete chatbot"}


# create a openain function to analyze a converstation extract answer of thses question
def analyze_conversation(chatbot, session_id=""):
    converstation = get_details_by_id(chatbot.bot_id, session_id)

    # openai
    questions = chatbot.question
    # call openai function
    answer = analyze_article(converstation, questions)
    # update analysis in database
    try:
        update_analyzation(chatbot.bot_id, session_id, answer)
    except:
        pass
    return answer


# fubction to analyze session
def analyze_session(chatbot, session_id=""):
    converstation = get_details_by_id(chatbot.bot_id, session_id)
    # call openai function
    answer = analyze_conversation_openai(converstation)
    # update analysis in database
    try:
        update_analytic(chatbot.bot_id, session_id, answer)
    except:
        pass
    return answer


# use scrape api to get data from url, argument is url
def get_data_from_url_scrape_api(url=""):
    urls = "https://scrap.duality.seethos.com/scrape-sitemap_links"
    headers = {
        'accept': 'application/json',
        'Content-Type': 'application/json',
    }
    print(url)
    data = {
        "url": url
    }
    # send data to url
    response = requests.post(urls, headers=headers, data=json.dumps(data))
    # Check the response
    if response.status_code == 200:
        print('Request successful')
        return response.json()['links']
    else:
        print(f'Request failed with status code: {response.status_code}')
        print(response.text)
        return {}


@app.task
def analyze_all_session(subdomain, chatbot):
    # get all session id
    session = get_all_session_by_subdomain(subdomain)
    # loop through session ids
    for session in session:
        # call analyze session function
        analyze_session(session_id=session['_id'], chatbot=chatbot)
        analyze_conversation(session_id=session['_id'], chatbot=chatbot)
    return {"message": "success"}
